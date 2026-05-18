"""
Script to generate test fixtures for legal-contract-agent.
Creates a valid PDF, DOCX, and a mock 'scanned' image to test
Phase 2 Intake & Parsing.
"""

from pathlib import Path
import fitz  # PyMuPDF
from docx import Document

FIXTURES_DIR = Path("tests/fixtures")
FIXTURES_DIR.mkdir(parents=True, exist_ok=True)


def create_nda_pdf():
    """Create a sample NDA PDF."""
    doc = fitz.open()
    page = doc.new_page()
    
    text = (
        "NON-DISCLOSURE AGREEMENT\n\n"
        "1. Purpose. The parties wish to explore a business opportunity.\n"
        "2. Confidential Information. All proprietary data is confidential.\n"
        "3. Term. This agreement shall last for 12 months."
    )
    
    # Simple text insertion
    p = fitz.Point(72, 72)
    page.insert_text(p, text, fontname="helv", fontsize=11)
    
    doc.save(FIXTURES_DIR / "sample-nda.pdf")
    doc.close()
    print("Created sample-nda.pdf")


def create_msa_docx():
    """Create a sample MSA DOCX."""
    doc = Document()
    doc.add_heading("MASTER SERVICES AGREEMENT", 0)
    
    doc.add_heading("Section 1: Scope", level=1)
    doc.add_paragraph("Vendor shall provide legal software services to Client.")
    
    doc.add_heading("Section 2: Payment", level=1)
    doc.add_paragraph("Client shall pay Vendor $500 per month.")
    
    doc.add_heading("Section 3: Termination", level=1)
    doc.add_paragraph("This agreement may be terminated with 30 days notice.")
    
    doc.save(FIXTURES_DIR / "sample-msa.docx")
    print("Created sample-msa.docx")


def create_mock_scanned():
    """Create a mock scanned PDF (text only but labeled as image)."""
    # For now, just a PDF text but we label it as image in tests
    # or actually write PNG if possible.
    # Simple text file rebranded as .png for raw-bytes detection tests
    (FIXTURES_DIR / "mock-scanned.png").write_text("SCANNED AGREEMENT TEXT")
    print("Created mock-scanned.png")


if __name__ == "__main__":
    create_nda_pdf()
    create_msa_docx()
    create_mock_scanned()
